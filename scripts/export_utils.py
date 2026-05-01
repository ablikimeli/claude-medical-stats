#!/usr/bin/env python3
"""
Export utilities: Word (.docx), Excel (.xlsx), PDF+PNG figures
Medical Statistics Skill
"""
import os
import pandas as pd
import numpy as np
from datetime import datetime


def export_to_excel(data, filename, sheet_name="Results", title=None):
    """
    Export data to Excel (.xlsx).
    data can be: pandas DataFrame, or list of (name, DataFrame) tuples for multi-sheet.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("  ! openpyxl not installed. Installing...")
        import subprocess
        subprocess.check_call(["pip", "install", "openpyxl"])
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    # Remove default sheet
    wb.remove(wb.active)

    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    header_font = Font(bold=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    if isinstance(data, list):
        sheets = data
    else:
        sheets = [(sheet_name, data)]

    for idx, (name, df) in enumerate(sheets):
        ws = wb.create_sheet(title=name[:31])  # Excel max sheet name length

        row_offset = 1
        if title and idx == 0:
            ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=13)
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
            row_offset = 2

        # Write headers
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=row_offset, column=col_idx, value=str(col_name))
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')

        # Write data
        for row_idx, (_, row) in enumerate(df.iterrows(), row_offset + 1):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                if isinstance(value, float):
                    cell.alignment = Alignment(horizontal='center')
                    cell.number_format = '0.0000'
                elif isinstance(value, (int, np.integer)):
                    cell.alignment = Alignment(horizontal='center')
                else:
                    cell.alignment = Alignment(horizontal='center')

        # Auto-fit column widths
        for col_idx in range(1, len(df.columns) + 1):
            max_len = 0
            for row_idx in range(row_offset, ws.max_row + 1):
                cell_val = ws.cell(row=row_idx, column=col_idx).value
                if cell_val is not None:
                    max_len = max(max_len, len(str(cell_val)))
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_len + 4, 40)

    wb.save(filename)
    print(f"  ✔ Excel exported: {os.path.abspath(filename)}")
    return filename


def export_to_word(table_data, filename, title=None):
    """
    Export a DataFrame to Word (.docx) as a formatted table.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("  ! python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

    doc = Document()

    # Title
    if title:
        doc.add_heading(title, level=1)

    # Convert DataFrame to Word table
    df = table_data.copy()

    # Handle MultiIndex columns
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = ['_'.join(col).strip() for col in df.columns]

    # Add table
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(df.columns):
        header_cells[col_idx].text = str(col_name)
        for paragraph in header_cells[col_idx].paragraphs:
            paragraph.alignment = 1  # center
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
        # Shading
        shading_elm = header_cells[col_idx]._tc.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E1F2',
            qn('w:val'): 'clear'
        })
        shading_elm.append(shading)

    # Data rows
    for row_idx, (_, row) in enumerate(df.iterrows()):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx + 1].cells[col_idx]
            text = str(value) if not (isinstance(value, float) and np.isnan(value)) else ''
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'

    doc.save(filename)
    print(f"  ✔ Word exported: {os.path.abspath(filename)}")
    return filename


def save_plot_dual(fig, basename, width=8, height=6, dpi=300):
    """
    Save a matplotlib figure as both PNG and PDF.
    """
    png_file = f"{basename}.png"
    fig.savefig(png_file, dpi=dpi, bbox_inches="tight", width=width, height=height)
    print(f"  ✔ PNG saved: {os.path.abspath(png_file)}")

    pdf_file = f"{basename}.pdf"
    fig.savefig(pdf_file, format="pdf", bbox_inches="tight", width=width, height=height)
    print(f"  ✔ PDF saved: {os.path.abspath(pdf_file)}")

    return png_file, pdf_file


def timestamp():
    """Generate timestamp string for filenames."""
    return datetime.now().strftime("%Y%m%d_%H%M%S")


if __name__ == "__main__":
    print("✔ export_utils.py loaded")
